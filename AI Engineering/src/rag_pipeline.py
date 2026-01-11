"""
RAG Pipeline for AI Engineering
===============================

Document loading, chunking, retrieval, and RAG chain components.
"""

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Callable


# =============================================================================
# Document Loading
# =============================================================================


@dataclass
class Document:
    """A document with content and metadata."""

    content: str
    metadata: dict = field(default_factory=dict)
    source: str = ""


class DocumentLoader:
    """
    Load documents from various sources.

    Example
    -------
    >>> loader = DocumentLoader()
    >>> docs = loader.load_file("document.txt")
    """

    def load_file(self, path: str) -> list[Document]:
        """
        Load a single file.

        Supports: .txt, .md, .pdf, .docx
        """
        path = Path(path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        extension = path.suffix.lower()

        if extension in [".txt", ".md"]:
            return self._load_text(path)
        elif extension == ".pdf":
            return self._load_pdf(path)
        elif extension == ".docx":
            return self._load_docx(path)
        else:
            # Try as text
            return self._load_text(path)

    def load_directory(
        self,
        directory: str,
        extensions: list[str] = None,
    ) -> list[Document]:
        """Load all documents from a directory."""
        directory = Path(directory)
        extensions = extensions or [".txt", ".md", ".pdf", ".docx"]

        documents = []
        for path in directory.rglob("*"):
            if path.is_file() and path.suffix.lower() in extensions:
                try:
                    docs = self.load_file(str(path))
                    documents.extend(docs)
                except Exception as e:
                    print(f"Error loading {path}: {e}")

        print(f"Loaded {len(documents)} documents from {directory}")
        return documents

    def _load_text(self, path: Path) -> list[Document]:
        """Load a text file."""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        return [
            Document(
                content=content,
                metadata={"file_type": "text"},
                source=str(path),
            )
        ]

    def _load_pdf(self, path: Path) -> list[Document]:
        """Load a PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            documents = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    documents.append(
                        Document(
                            content=text,
                            metadata={"file_type": "pdf", "page": i + 1},
                            source=str(path),
                        )
                    )

            return documents
        except ImportError:
            raise ImportError("pypdf is required for PDF loading. Install with: pip install pypdf")

    def _load_docx(self, path: Path) -> list[Document]:
        """Load a Word document."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            content = "\n".join([para.text for para in doc.paragraphs])

            return [
                Document(
                    content=content,
                    metadata={"file_type": "docx"},
                    source=str(path),
                )
            ]
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")


# =============================================================================
# Chunking
# =============================================================================


class Chunker:
    """
    Split documents into chunks for embedding.

    Example
    -------
    >>> chunker = Chunker(chunk_size=500, overlap=50)
    >>> chunks = chunker.chunk(document)
    """

    def __init__(
        self,
        chunk_size: int = 500,
        overlap: int = 50,
        separator: str = "\n",
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.separator = separator

    def chunk(self, document: Document) -> list[Document]:
        """Split a document into chunks."""
        text = document.content
        chunks = []

        # Split by separator first
        sections = text.split(self.separator)

        current_chunk = ""
        for section in sections:
            # If adding this section exceeds chunk size
            if len(current_chunk) + len(section) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())

                # Start new chunk with overlap
                if self.overlap > 0 and current_chunk:
                    overlap_text = current_chunk[-self.overlap:]
                    current_chunk = overlap_text + self.separator + section
                else:
                    current_chunk = section
            else:
                current_chunk += self.separator + section if current_chunk else section

        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # Convert to Document objects
        return [
            Document(
                content=chunk,
                metadata={
                    **document.metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                },
                source=document.source,
            )
            for i, chunk in enumerate(chunks)
        ]

    def chunk_all(self, documents: list[Document]) -> list[Document]:
        """Chunk multiple documents."""
        all_chunks = []
        for doc in documents:
            chunks = self.chunk(doc)
            all_chunks.extend(chunks)

        print(f"Created {len(all_chunks)} chunks from {len(documents)} documents")
        return all_chunks


# =============================================================================
# Vector Store (using ChromaDB)
# =============================================================================


class VectorStore:
    """
    Vector store using ChromaDB.

    Example
    -------
    >>> store = VectorStore(collection_name="my_docs")
    >>> store.add_documents(chunks)
    >>> results = store.search("query", k=5)
    """

    def __init__(
        self,
        collection_name: str = "documents",
        persist_directory: Optional[str] = None,
    ):
        try:
            import chromadb
        except ImportError:
            raise ImportError("chromadb is required. Install with: pip install chromadb")

        if persist_directory:
            self.client = chromadb.PersistentClient(path=persist_directory)
        else:
            self.client = chromadb.Client()

        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

    def add_documents(
        self,
        documents: list[Document],
        embedding_function: Optional[Callable] = None,
    ) -> None:
        """Add documents to the store."""
        texts = [doc.content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        ids = [f"doc_{i}" for i in range(len(documents))]

        if embedding_function:
            embeddings = embedding_function(texts)
            self.collection.add(
                documents=texts,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids,
            )
        else:
            # Use ChromaDB's default embedding
            self.collection.add(
                documents=texts,
                metadatas=metadatas,
                ids=ids,
            )

        print(f"Added {len(documents)} documents to vector store")

    def search(
        self,
        query: str,
        k: int = 5,
        filter: Optional[dict] = None,
    ) -> list[dict]:
        """Search for similar documents."""
        results = self.collection.query(
            query_texts=[query],
            n_results=k,
            where=filter,
        )

        # Format results
        formatted = []
        for i in range(len(results["documents"][0])):
            formatted.append({
                "content": results["documents"][0][i],
                "metadata": results["metadatas"][0][i] if results["metadatas"] else {},
                "distance": results["distances"][0][i] if results["distances"] else 0,
            })

        return formatted

    def count(self) -> int:
        """Get number of documents in store."""
        return self.collection.count()


# =============================================================================
# RAG Pipeline
# =============================================================================


class RAGPipeline:
    """
    Complete RAG pipeline.

    Example
    -------
    >>> rag = RAGPipeline(llm_client)
    >>> rag.load_documents("./docs")
    >>> answer = rag.query("What is the main topic?")
    """

    def __init__(
        self,
        llm_client,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        collection_name: str = "rag_docs",
    ):
        self.llm = llm_client
        self.loader = DocumentLoader()
        self.chunker = Chunker(chunk_size=chunk_size, overlap=chunk_overlap)
        self.vector_store = VectorStore(collection_name=collection_name)

    def load_documents(self, path: str) -> int:
        """Load and index documents from a file or directory."""
        path = Path(path)

        if path.is_file():
            documents = self.loader.load_file(str(path))
        else:
            documents = self.loader.load_directory(str(path))

        # Chunk documents
        chunks = self.chunker.chunk_all(documents)

        # Add to vector store
        self.vector_store.add_documents(chunks)

        return len(chunks)

    def query(
        self,
        question: str,
        k: int = 5,
        system_prompt: Optional[str] = None,
    ) -> dict:
        """
        Answer a question using RAG.

        Returns
        -------
        dict
            {answer, sources, context}
        """
        # Retrieve relevant documents
        results = self.vector_store.search(question, k=k)

        # Build context
        context = "\n\n---\n\n".join([r["content"] for r in results])

        # Generate answer
        prompt = f"""Use the following context to answer the question.
If you cannot answer from the context, say so.

Context:
{context}

Question: {question}

Answer:"""

        answer = self.llm.chat(prompt, system=system_prompt)

        return {
            "answer": answer,
            "sources": results,
            "context": context,
        }

    def get_stats(self) -> dict:
        """Get pipeline statistics."""
        return {
            "documents": self.vector_store.count(),
            "llm_stats": self.llm.get_stats().summary() if hasattr(self.llm, 'get_stats') else "N/A",
        }
